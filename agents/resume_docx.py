"""
Render a tailored resume (markdown) to a .docx file.

The resume markdown produced by the pipeline uses a small, predictable subset
of markdown:

    # NAME                  -> document title
    ## Role / subtitle      -> subtitle under the name
    plain lines             -> contact lines / body paragraphs
    ---                     -> section divider (rendered as a horizontal rule)
    # SECTION               -> section heading (e.g. EMPLOYMENT HISTORY)
    ## Job | Company        -> job / entry heading
    **bold text**           -> bold paragraph (e.g. dates)
    - bullet                -> bullet list item

This converter targets that subset rather than being a general markdown engine.
Inline `**bold**` spans inside any line are rendered as bold runs.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ACCENT = RGBColor(0x1A, 0x3C, 0x5E)


def _add_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph`, rendering **...** spans as bold runs."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_divider(doc) -> None:
    """Add a thin empty paragraph to act as visual spacing between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def markdown_to_docx(markdown: str, out_path: str | Path) -> Path:
    """Convert resume `markdown` to a .docx written at `out_path`."""
    out_path = Path(out_path)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    first_h1_seen = False
    subtitle_pending = False  # next ## directly follows the name -> center it

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.strip() == "---":
            _add_divider(doc)
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if not first_h1_seen:
                # Top-level name -> document title
                first_h1_seen = True
                subtitle_pending = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(22)
                run.font.color.rgb = _ACCENT
            else:
                # Section heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = _ACCENT
                subtitle_pending = False
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            # Subtitle directly under the name is centered.
            if subtitle_pending:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_pending = False
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            continue

        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, text)
            continue

        # Plain paragraph (contact lines, summary text, bold dates, etc.)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_runs(p, line)

    doc.save(out_path)
    return out_path
