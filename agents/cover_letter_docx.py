"""
Render a cover letter (markdown/plain prose) to a .docx file.

The cover letter produced by the pipeline is plain prose: a date line, a
salutation, body paragraphs, and a sign-off block, one item per line with
blank lines between paragraphs. Inline `**bold**` spans are rendered as bold
runs (rare, but supported for consistency with the resume converter).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from agents.resume_docx import _add_runs


def markdown_to_docx(markdown: str, out_path: str | Path) -> Path:
    """Convert cover letter `markdown` to a .docx written at `out_path`."""
    out_path = Path(out_path)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _add_runs(p, line)

    doc.save(out_path)
    return out_path
